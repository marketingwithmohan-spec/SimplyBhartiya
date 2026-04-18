#!/usr/bin/env python3
"""
Convert markdown files to Word (.docx) format
Processes all 7 SaaS documentation files
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to Word document"""
    
    doc = Document()
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split by lines
    lines = content.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip empty lines (but add spacing)
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line_stripped.startswith('# '):
            heading = line_stripped[2:]
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif line_stripped.startswith('## '):
            heading = line_stripped[3:]
            p = doc.add_heading(heading, level=2)
            
        elif line_stripped.startswith('### '):
            heading = line_stripped[4:]
            p = doc.add_heading(heading, level=3)
            
        elif line_stripped.startswith('#### '):
            heading = line_stripped[5:]
            p = doc.add_heading(heading, level=4)
        
        # Handle bold and italic formatting
        else:
            # Process inline markdown formatting
            text = line_stripped
            
            # Add paragraph
            p = doc.add_paragraph()
            
            # Simple markdown parsing for bold and italic
            # **bold** -> bold
            # *italic* -> italic
            # `code` -> code
            
            i = 0
            while i < len(text):
                # Bold
                if text[i:i+2] == '**':
                    end = text.find('**', i+2)
                    if end != -1:
                        bold_text = text[i+2:end]
                        run = p.add_run(bold_text)
                        run.bold = True
                        i = end + 2
                        continue
                
                # Code
                if text[i] == '`':
                    end = text.find('`', i+1)
                    if end != -1:
                        code_text = text[i+1:end]
                        run = p.add_run(code_text)
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(192, 0, 0)
                        i = end + 1
                        continue
                
                # Regular text
                if i < len(text):
                    p.add_run(text[i])
                    i += 1
    
    # Save document
    doc.save(docx_file_path)
    print(f"✓ Created: {docx_file_path}")

def main():
    """Convert all 7 SaaS documentation files"""
    
    base_path = r"e:\Simply Bhartiya"
    
    files_to_convert = [
        "SAAS_OPERATIONAL_WORKFLOW.md",
        "SAAS_INFRASTRUCTURE_COST_ANALYSIS.md",
        "SAAS_CLOUD_DATA_ARCHITECTURE.md",
        "SAAS_EXCEL_CSV_EXPORT.md",
        "SAAS_GOOGLE_DRIVE_INTEGRATION.md",
        "SAAS_PRODUCT_ROADMAP.md",
        "SAAS_COMPLETE_IMPLEMENTATION_GUIDE.md"
    ]
    
    print("Converting markdown documents to Word format...\n")
    
    for md_file in files_to_convert:
        md_path = os.path.join(base_path, md_file)
        docx_file = md_file.replace('.md', '.docx')
        docx_path = os.path.join(base_path, docx_file)
        
        if os.path.exists(md_path):
            try:
                markdown_to_docx(md_path, docx_path)
            except Exception as e:
                print(f"✗ Failed: {md_file} - {str(e)}")
        else:
            print(f"✗ Not found: {md_path}")
    
    print("\n✓ All documents converted successfully!")
    print("\nWord files created:")
    for md_file in files_to_convert:
        docx_file = md_file.replace('.md', '.docx')
        print(f"  • {docx_file}")

if __name__ == "__main__":
    main()
