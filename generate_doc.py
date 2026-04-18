from docx import Document
from docx.shared import Pt

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(12)

title = doc.add_paragraph('Deployment Recommendations for Simply Bhartiya', style='Title')

lines = [
    '1. Use GitHub or a similar code repository to store your project code.',
    '2. Host the database on MongoDB Atlas free tier for small business use.',
    '3. Deploy the backend to a simple cloud service like Render.com or Railway.app.',
    '4. Deploy the frontend as a static website on Vercel, Netlify, or Cloudflare Pages.',
    '5. Set environment variables in production: MONGODB_URI, PORT, NODE_ENV.',
    '6. Update frontend API_BASE_URL to the live backend URL after deployment.',
    '7. Secure the app with HTTPS, strong database credentials, and do not publish secrets.',
    '8. Use Stripe or Razorpay later to add subscription billing and user control.',
]

doc.add_paragraph('')
for line in lines:
    doc.add_paragraph(line)

doc.add_paragraph('')
doc.add_paragraph('Final Recommendation', style='Heading 2')
doc.add_paragraph('For a small oil-extracting company, start with the free tiers of Atlas, Render, and Vercel. Once the app is stable and traffic grows, move to paid plans. Add authentication and subscription billing only after the app is live and working well.')

doc.save(r'e:/Simply Bhartiya/Deployment-Recommendation.docx')
print('Created e:/Simply Bhartiya/Deployment-Recommendation.docx')
